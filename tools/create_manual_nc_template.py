from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "material-video" / "antes-evidran" / "Plantilla-NC-manual-antes-Evidran.docx"

# Preset: compact_reference_guide.
# Named overrides: FORM_MASTHEAD, FORM_SECTION_BAND, FORM_LABEL,
# FORM_PLACEHOLDER and FORM_TABLE_DENSITY.
NAVY = "1F4E78"
BLUE_GRAY = "D9E2F3"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9D9D9"
BORDER = "8A9AAA"
INK = RGBColor(31, 31, 31)
MUTED = RGBColor(100, 108, 117)
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_run(run, size=9.5, bold=False, italic=False, color=INK, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def format_paragraph(paragraph, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT, keep=False):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def clear_cell(cell):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    format_paragraph(paragraph)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return paragraph


def add_text(cell, text, size=9.5, bold=False, italic=False, color=INK,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.0):
    paragraph = clear_cell(cell)
    format_paragraph(paragraph, before=before, after=after, line=line, align=align)
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_label(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    set_cell_shading(cell, LIGHT_BLUE)
    return add_text(cell, text.upper(), size=8.5, bold=True, color=RGBColor(31, 78, 120), align=align)


def add_placeholder(cell, text="[ Haga clic aquí y escriba... ]", min_height=0.46):
    add_text(cell, text, size=9.5, italic=True, color=MUTED, line=1.15)
    row = cell._tc.getparent()
    tr_pr = row.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(int(min_height * 1440)))
    tr_height.set(qn("w:hRule"), "atLeast")


def add_table(doc, rows, widths, border_color=BORDER, header_repeat=False):
    table = doc.add_table(rows=rows, cols=len(widths))
    set_table_geometry(table, widths)
    set_table_borders(table, border_color)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if header_repeat and table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
    return table


def add_doc_masthead(doc, page_label, full=False):
    table = add_table(doc, 1, [1700, 5100, 2560], border_color=NAVY)
    add_text(table.cell(0, 0), "[ LOGOTIPO\nDE LA EMPRESA ]", size=9, bold=True,
             color=RGBColor(90, 90, 90), align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    center = table.cell(0, 1)
    set_cell_shading(center, NAVY)
    p = clear_cell(center)
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    r = p.add_run("INFORME DE NO CONFORMIDAD" if full else page_label.upper())
    set_run(r, size=14 if full else 11, bold=True, color=WHITE)
    if full:
        p2 = center.add_paragraph()
        format_paragraph(p2, before=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        r2 = p2.add_run("Sistema de Gestión de la Calidad · UNE-EN ISO 9001:2015 · 10.2")
        set_run(r2, size=8, color=WHITE)
    right = table.cell(0, 2)
    p = clear_cell(right)
    format_paragraph(p, line=1.05)
    for idx, line in enumerate(("Código: F-CAL-07", "Edición: 03", "Fecha: __ / __ / ____", "Página: " + page_label)):
        target = p if idx == 0 else right.add_paragraph()
        format_paragraph(target, after=0, line=1.0)
        run = target.add_run(line)
        set_run(run, size=7.8, bold=(idx == 0), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_section_band(doc, number, title, note=None):
    table = add_table(doc, 1, [900, 8460], border_color=NAVY)
    set_cell_shading(table.cell(0, 0), NAVY)
    add_text(table.cell(0, 0), f"{number:02d}", size=12, bold=True, color=WHITE,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(0, 1), BLUE_GRAY)
    p = clear_cell(table.cell(0, 1))
    format_paragraph(p, line=1.0)
    r = p.add_run(title.upper())
    set_run(r, size=10.5, bold=True, color=RGBColor(31, 78, 120))
    if note:
        r2 = p.add_run(f"  ·  {note}")
        set_run(r2, size=8.2, italic=True, color=MUTED)


def add_spacer(doc, points=4):
    p = doc.add_paragraph()
    format_paragraph(p, after=points)


def page_break(doc):
    doc.add_page_break()


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, RGBColor(46, 116, 181), 18, 10),
        ("Heading 2", 13, RGBColor(46, 116, 181), 14, 7),
        ("Heading 3", 12, RGBColor(31, 77, 120), 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("F-CAL-07 · INFORME DE NO CONFORMIDAD · Copia no controlada   |   Página ")
    set_run(r, size=8, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.extend([fld_char, instr_text, fld_sep, fallback, fld_end])
    r2 = p.add_run(" de 4")
    set_run(r2, size=8, color=MUTED)


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    # Compact four-page form: the dense layout is intentional for the
    # "before Evidran" recording, while keeping every block on its page.
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.24)
    add_footer(section)

    doc.core_properties.title = "Plantilla manual de informe de no conformidad"
    doc.core_properties.subject = "Material de demostración - antes de Evidran"
    doc.core_properties.author = "Evidran"
    doc.core_properties.comments = "Plantilla ficticia para grabar un proceso manual y tedioso."

    # PAGE 1 · Registro y contención
    add_doc_masthead(doc, "1 de 4", full=True)
    add_section_band(doc, 1, "Identificación y descripción", "Todos los campos son obligatorios")
    meta = add_table(doc, 4, [1500, 3180, 1500, 3180])
    labels = [
        ("N.º de NC", "[ NC-AAAA-NNN ]", "Fecha de detección", "[ dd/mm/aaaa ]"),
        ("Detectado por", "[ Nombre y puesto ]", "Origen", "[ Auditoría / cliente / interna ]"),
        ("Proceso / área", "[ Proceso afectado ]", "Producto / lote", "[ Referencia, lote y cantidad ]"),
        ("Norma / requisito", "[ Norma y apartado ]", "Clasificación", "☐ Mayor   ☐ Menor   ☐ Observación"),
    ]
    for r_idx, row_values in enumerate(labels):
        add_label(meta.cell(r_idx, 0), row_values[0])
        add_placeholder(meta.cell(r_idx, 1), row_values[1], 0.36)
        add_label(meta.cell(r_idx, 2), row_values[2])
        add_placeholder(meta.cell(r_idx, 3), row_values[3], 0.36)

    add_spacer(doc, 3)
    desc = add_table(doc, 4, [9360])
    add_label(desc.cell(0, 0), "Descripción objetiva de la no conformidad")
    add_placeholder(desc.cell(1, 0), "[ Describa qué ocurrió, dónde, cuándo y cómo se detectó. Evite opiniones. ]", 0.72)
    add_label(desc.cell(2, 0), "Alcance conocido y posible extensión")
    add_placeholder(desc.cell(3, 0), "[ Indique cantidades, lotes, procesos, turnos, documentos y clientes que podrían estar afectados. ]", 0.70)

    add_spacer(doc, 3)
    correction = add_table(doc, 4, [9360])
    add_label(correction.cell(0, 0), "Corrección o contención inmediata")
    add_placeholder(correction.cell(1, 0), "[ Detalle la segregación, retrabajo, bloqueo, aviso o corrección aplicada. ]", 0.62)
    add_label(correction.cell(2, 0), "Responsable, fecha y comprobación de la contención")
    add_placeholder(correction.cell(3, 0), "[ Responsable: __________   Fecha: ___/___/_____   Resultado: ______________________________ ]", 0.43)

    add_spacer(doc, 3)
    ev = add_table(doc, 4, [750, 4730, 1800, 2080], header_repeat=True)
    for c, text in enumerate(("N.º", "Evidencia / archivo", "Adjuntado por", "Fecha")):
        add_label(ev.cell(0, c), text, WD_ALIGN_PARAGRAPH.CENTER)
    for r_idx in range(1, 4):
        add_text(ev.cell(r_idx, 0), str(r_idx), size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_placeholder(ev.cell(r_idx, 1), "[ nombre-del-archivo.ext ]", 0.34)
        add_placeholder(ev.cell(r_idx, 2), "[ Nombre ]", 0.34)
        add_placeholder(ev.cell(r_idx, 3), "[ dd/mm/aaaa ]", 0.34)

    # PAGE 2 · Analysis
    page_break(doc)
    add_doc_masthead(doc, "2 de 4")
    add_section_band(doc, 2, "Análisis de causa raíz", "Complete la secuencia y justifique la causa final")
    method = add_table(doc, 2, [2200, 7160])
    add_label(method.cell(0, 0), "Método utilizado")
    add_text(method.cell(0, 1), "☐ 5 porqués   ☐ Ishikawa   ☐ Pareto   ☐ Otro: ____________________", size=9)
    add_label(method.cell(1, 0), "Equipo investigador")
    add_placeholder(method.cell(1, 1), "[ Nombres, cargos y áreas participantes ]", 0.38)

    add_spacer(doc, 3)
    whys = add_table(doc, 6, [820, 2620, 5920], header_repeat=True)
    for c, text in enumerate(("Paso", "Pregunta", "Respuesta / evidencia")):
        add_label(whys.cell(0, c), text, WD_ALIGN_PARAGRAPH.CENTER)
    for idx in range(1, 6):
        add_text(whys.cell(idx, 0), str(idx), size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(whys.cell(idx, 1), f"¿Por qué ocurrió?" if idx == 1 else f"¿Por qué ocurrió lo anterior?", size=8.8)
        add_placeholder(whys.cell(idx, 2), "[ Escriba la respuesta y cite el registro que la sostiene. ]", 0.45)

    add_spacer(doc, 3)
    root = add_table(doc, 6, [9360])
    add_label(root.cell(0, 0), "Causa raíz de ocurrencia")
    add_placeholder(root.cell(1, 0), "[ Formule la causa de sistema, no un síntoma ni 'error humano'. ]", 0.40)
    add_label(root.cell(2, 0), "Por qué no se detectó antes")
    add_placeholder(root.cell(3, 0), "[ Explique qué control faltó, falló o no estaba definido. ]", 0.38)
    add_label(root.cell(4, 0), "Comprobación de extensión")
    add_placeholder(root.cell(5, 0), "[ Revise otros lotes, equipos, turnos, procesos, documentos e informes emitidos. Indique resultado. ]", 0.40)

    add_spacer(doc, 3)
    approval = add_table(doc, 2, [2200, 7160])
    add_label(approval.cell(0, 0), "Conclusión del análisis")
    add_text(approval.cell(0, 1), "☐ Causa confirmada   ☐ Hipótesis pendiente   ☐ Requiere investigación adicional", size=8.8)
    add_label(approval.cell(1, 0), "Revisado por / fecha")
    add_placeholder(approval.cell(1, 1), "[ Nombre: ____________________   Firma: ________________   Fecha: ___/___/_____ ]", 0.32)

    # PAGE 3 · Actions
    page_break(doc)
    add_doc_masthead(doc, "3 de 4")
    add_section_band(doc, 3, "Plan de acciones", "Asigne responsable, plazo y evidencia para cada medida")
    decision = add_table(doc, 3, [2200, 7160])
    add_label(decision.cell(0, 0), "¿Requiere acción correctiva?")
    add_text(decision.cell(0, 1), "☐ Sí   ☐ No   ☐ Pendiente de decisión", size=9)
    add_label(decision.cell(1, 0), "Justificación")
    add_placeholder(decision.cell(1, 1), "[ Explique la decisión considerando recurrencia, gravedad y alcance. ]", 0.36)
    add_label(decision.cell(2, 0), "Prioridad")
    add_text(decision.cell(2, 1), "☐ Alta   ☐ Media   ☐ Baja      Riesgo asociado: [ __________ ]", size=9)

    add_spacer(doc, 3)
    actions = add_table(doc, 7, [620, 3370, 1150, 1120, 1030, 2070], header_repeat=True)
    headers = ("N.º", "Acción / entregable", "Responsable", "Fecha límite", "Estado", "Evidencia prevista")
    for c, text in enumerate(headers):
        add_label(actions.cell(0, c), text, WD_ALIGN_PARAGRAPH.CENTER)
    for idx in range(1, 7):
        add_text(actions.cell(idx, 0), str(idx), size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_placeholder(actions.cell(idx, 1), "[ Describa la acción ]", 0.42)
        add_placeholder(actions.cell(idx, 2), "[ Nombre ]", 0.42)
        add_placeholder(actions.cell(idx, 3), "[ dd/mm/aa ]", 0.42)
        add_placeholder(actions.cell(idx, 4), "[ Estado ]", 0.42)
        add_placeholder(actions.cell(idx, 5), "[ Registro / archivo ]", 0.42)

    add_spacer(doc, 3)
    change = add_table(doc, 6, [2200, 7160])
    fields = (
        ("Documentos a modificar", "[ Procedimiento, instrucción, pauta, formato o plan afectado ]"),
        ("Formación / comunicación", "[ Personas afectadas, contenido y fecha prevista ]"),
        ("Recursos necesarios", "[ Tiempo, presupuesto, equipos o soporte externo ]"),
        ("Riesgos de la acción", "[ Posibles efectos no deseados y controles ]"),
        ("Seguimiento intermedio", "[ Reunión / fecha / responsable / resultado esperado ]"),
        ("Aprobación del plan", "[ Nombre: ____________________   Firma: ________________   Fecha: ___/___/_____ ]"),
    )
    for idx, (label, placeholder) in enumerate(fields):
        add_label(change.cell(idx, 0), label)
        add_placeholder(change.cell(idx, 1), placeholder, 0.30 if idx != 3 else 0.34)

    # PAGE 4 · Effectiveness and closure
    page_break(doc)
    add_doc_masthead(doc, "4 de 4")
    add_section_band(doc, 4, "Verificación de eficacia y cierre", "No cerrar hasta disponer de evidencia objetiva")
    verify = add_table(doc, 6, [2200, 7160])
    fields = (
        ("Método de verificación", "[ Auditoría, revisión de registros, medición, muestreo, ausencia de recurrencia... ]"),
        ("Indicador / criterio", "[ Qué se medirá y con qué fórmula o regla de aceptación ]"),
        ("Muestra / periodo", "[ Cantidad de lotes, semanas, registros o personas a revisar ]"),
        ("Resultado objetivo", "[ Meta cuantitativa o condición que debe cumplirse ]"),
        ("Fecha de verificación", "[ dd/mm/aaaa ]"),
        ("Responsable de verificar", "[ Nombre y cargo; preferiblemente distinto del ejecutor ]"),
    )
    for idx, (label, placeholder) in enumerate(fields):
        add_label(verify.cell(idx, 0), label)
        add_placeholder(verify.cell(idx, 1), placeholder, 0.30)

    add_spacer(doc, 3)
    result = add_table(doc, 6, [9360])
    add_label(result.cell(0, 0), "Resultado obtenido y evidencias revisadas")
    add_placeholder(result.cell(1, 0), "[ Registre datos, fechas, tamaño de muestra y referencias de evidencias. ]", 0.50)
    add_label(result.cell(2, 0), "Conclusión sobre la eficacia")
    add_text(result.cell(3, 0), "☐ Eficaz   ☐ Parcialmente eficaz   ☐ No eficaz   ☐ No verificable todavía", size=9)
    add_label(result.cell(4, 0), "Acciones si no es eficaz")
    add_placeholder(result.cell(5, 0), "[ Reabrir expediente / nueva acción / ampliar análisis / nueva fecha de verificación. ]", 0.34)

    add_spacer(doc, 3)
    sign = add_table(doc, 4, [1800, 2520, 2520, 2520], header_repeat=True)
    for c, text in enumerate(("Rol", "Nombre y cargo", "Firma", "Fecha")):
        add_label(sign.cell(0, c), text, WD_ALIGN_PARAGRAPH.CENTER)
    roles = ("Elaborado por", "Revisado por", "Aprobado por")
    for idx, role in enumerate(roles, start=1):
        add_label(sign.cell(idx, 0), role)
        add_placeholder(sign.cell(idx, 1), "[ Nombre ]", 0.38)
        add_placeholder(sign.cell(idx, 2), "[ Firma ]", 0.38)
        add_placeholder(sign.cell(idx, 3), "[ dd/mm/aaaa ]", 0.38)

    add_spacer(doc, 3)
    closing = add_table(doc, 3, [2200, 7160])
    add_label(closing.cell(0, 0), "Estado final")
    add_text(closing.cell(0, 1), "☐ Cerrada   ☐ Reabierta   ☐ Pendiente de eficacia", size=9)
    add_label(closing.cell(1, 0), "Fecha de cierre")
    add_placeholder(closing.cell(1, 1), "[ dd/mm/aaaa ]", 0.28)
    add_label(closing.cell(2, 0), "Distribución / archivo")
    add_placeholder(closing.cell(2, 1), "[ Carpeta, destinatarios, versión PDF y ubicación de evidencias ]", 0.30)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
